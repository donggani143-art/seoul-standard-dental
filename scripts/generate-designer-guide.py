# -*- coding: utf-8 -*-
"""
METALINK CMS 웹 퍼블리싱 디자이너 가이드 .docx 생성
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# 기본 폰트 설정
style = doc.styles['Normal']
font = style.font
font.name = 'Pretendard'
font.size = Pt(10.5)
rPr = style.element.rPr
rFonts = rPr.find(qn('w:rFonts')) if rPr is not None else None
if rFonts is not None:
    rFonts.set(qn('w:eastAsia'), 'Pretendard')


def add_title(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Pretendard'
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), 'Pretendard')
        rPr.append(rFonts)
    return h


def add_para(text, bold=False, color=None, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Pretendard'
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'Pretendard')
    rPr.append(rFonts)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = RGBColor(*color)
    if size:
        run.font.size = Pt(size)
    return p


def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.3)
    # 회색 배경 음영
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'F4F4F4')
    pPr.append(shd)
    return p


def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Pretendard'
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), 'Pretendard')
        rPr.append(rFonts)
    return p


def add_warning(text):
    p = doc.add_paragraph()
    run = p.add_run('⚠ ' + text)
    run.font.name = 'Pretendard'
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    run.bold = True
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'Pretendard')
    rPr.append(rFonts)
    return p


# ============= 표지 =============
title = doc.add_heading('METALINK CMS', 0)
for run in title.runs:
    run.font.name = 'Pretendard'
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'Pretendard')
    rPr.append(rFonts)
    run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

subtitle = add_para('웹 퍼블리싱 디자이너 작업 가이드', size=18, bold=True)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_para('').paragraph_format.space_after = Pt(40)

# ============= 1. 시작하기 =============
add_title('1. 시작하기', 1)

add_title('1.1 로그인', 2)
add_para('관리자 페이지에 접속하여 로그인합니다.')
add_code('https://metalink3.mycafe24.com/admin')
add_para('병원 관리자 계정 정보는 플랫폼 운영자(슈퍼관리자)로부터 발급받습니다.')

add_title('1.2 작업 환경', 2)
add_bullet('좌측 사이드바: 모든 관리 메뉴')
add_bullet('사이트 페이지: 페이지 빌더로 직접 코딩')
add_bullet('헤더·푸터: 공통 영역 편집')
add_bullet('SEO·병원 정보: 메타태그/위치/공통 코드 설정')
add_bullet('도메인 관리: 실제 도메인 연결')

# ============= 2. 페이지 빌더 =============
add_title('2. 페이지 빌더 사용법', 1)

add_title('2.1 구조', 2)
add_para('페이지 빌더는 3개의 탭으로 구성됩니다.', bold=True)
add_bullet('HTML: 페이지 구조 마크업')
add_bullet('CSS: 해당 페이지 전용 스타일')
add_bullet('JS: 해당 페이지 전용 JavaScript')

add_warning('JS 탭에는 <script> 태그를 쓰지 마세요. 순수 JavaScript 코드만 입력합니다. (자동 제거되긴 하지만 비추천)')

add_title('2.2 단축키', 2)
add_bullet('Ctrl + S: 저장')
add_bullet('Ctrl + F: 코드 내 검색')

add_title('2.3 미리보기 모드', 2)
add_bullet('분할 (기본): 좌측 코드, 우측 미리보기')
add_bullet('에디터: 코드만')
add_bullet('미리보기: 결과만')
add_bullet('데스크탑/모바일: 우측 디바이스 토글로 반응형 확인')

# ============= 3. 페이지 작성 규칙 =============
add_title('3. 페이지 작성 규칙', 1)

add_title('3.1 메인 페이지 (slug: main)', 2)
add_para('URL: /')
add_para('필수: 기본 제공되며 삭제 불가. 사이트의 첫 화면.')

add_title('3.2 공통 헤더/푸터 (slug: _header, _footer)', 2)
add_para('헤더·푸터 메뉴에서 편집. 모든 페이지 상하단에 자동 삽입됩니다.')
add_warning('헤더에 모바일 햄버거 메뉴를 만들 때 주의: PC용 .depth2와 모바일 .depth2 클래스명이 충돌하지 않게 작성하세요.')
add_para('해결책 예시:', bold=True)
add_code('''/* PC 드롭다운 */
.pc-nav .depth2 { position: absolute; opacity: 0; }

/* 모바일 아코디언 (PC 스타일 명시적 리셋) */
.m-accordion .depth2 {
  position: static;
  opacity: 1;
  visibility: visible;
  pointer-events: auto;
}''')

add_title('3.3 커스텀 페이지', 2)
add_para('사이트 페이지 → "+ 페이지 추가" → 기본/커스텀 선택')
add_bullet('기본: 미리 정의된 페이지(병원소개, 진료과목 등)에서 선택')
add_bullet('커스텀: 직접 slug와 제목 입력. URL은 /{slug} 자동 생성')

add_para('예시:', bold=True)
add_code('slug: about-equipment\n→ URL: https://yoursite.com/about-equipment')

# ============= 4. 블록 스니펫 =============
add_title('4. 블록 스니펫 활용', 1)

add_para('빌더 좌측 사이드바 > 블록 탭에서 한 번 클릭으로 HTML/CSS/JS를 삽입할 수 있습니다.')

add_title('4.1 제공 블록 목록', 2)
blocks = [
    ('인트로 애니메이션', '메인 페이지 진입 시 타이핑 + 브랜드 표시'),
    ('히어로 섹션', '큰 제목 + 부제 + CTA 버튼'),
    ('CTA 배너', '행동 유도 배너'),
    ('카드 3단', '서비스/특징 3개 나열'),
    ('텍스트 + 이미지', '좌우 2단 레이아웃'),
    ('인용/후기', '환자 후기 카드'),
    ('버튼', '단일 CTA 버튼'),
    ('구분선', '섹션 간 구분'),
    ('지도 임베드', 'iframe 지도'),
    ('퀵메뉴 (SNS/예약)', '우하단 플로팅 SNS 버튼 (헤더 또는 푸터에 삽입 권장)'),
    ('상담 신청 폼', '/api/consult 연동 완전한 폼'),
    ('환자 로그인 폼', '/api/auth/login 연동'),
    ('환자 회원가입 폼', '/api/auth/register 연동'),
    ('게시판 목록', 'API에서 게시글 자동 로드'),
    ('게시글 상세', 'URL ID로 게시글 표시'),
]
for name, desc in blocks:
    add_bullet('{}: {}'.format(name, desc))

add_title('4.2 게시판 블록 사용법', 2)
add_para('블록 삽입 후 boardType만 변경하세요.', bold=True)
add_code('''var boardType = 'notice';   // notice | event | board
var boardSlug = '';          // board일 때만 slug 입력''')
add_warning("boardType은 경로(community/notice)가 아닌 타입명(notice)만 입력!")

# ============= 5. 게시판 =============
add_title('5. 게시판 시스템', 1)

add_title('5.1 게시판 종류', 2)
add_bullet('공지사항 (notice): 시스템 내장')
add_bullet('이벤트 (event): 시스템 내장')
add_bullet('커스텀 게시판 (board): 게시판 관리에서 그룹 생성 후 사용')

add_title('5.2 게시판 목록 페이지 디자인', 2)
add_para('페이지 빌더에서 다음 슬러그로 디자인:')
add_bullet('notice: 공지사항 목록')
add_bullet('event: 이벤트 목록')
add_bullet('board-{slug}: 커스텀 게시판 목록 (예: board-free)')

add_title('5.3 게시글 상세 페이지 (템플릿 변수)', 2)
add_para('상세 페이지는 슬러그 끝에 -detail을 붙이고 템플릿 변수를 사용합니다.')
add_bullet('notice-detail: 공지 상세')
add_bullet('event-detail: 이벤트 상세')
add_bullet('board-{slug}-detail: 커스텀 상세')

add_para('사용 가능한 변수:', bold=True)
add_code('''{{post.title}}      → 게시글 제목
{{post.content}}    → 게시글 본문 (HTML)
{{post.date}}       → 등록일 (2026. 4. 8.)
{{post.boardLabel}} → 게시판 이름 (공지사항, 이벤트 등)
{{post.listHref}}   → 목록 페이지 링크
{{post.listLabel}}  → 목록 버튼 텍스트''')

add_para('예시:', bold=True)
add_code('''<article>
  <h1>{{post.title}}</h1>
  <p>{{post.date}} · {{post.boardLabel}}</p>
  <div>{{post.content}}</div>
  <a href="{{post.listHref}}">{{post.listLabel}}</a>
</article>''')

# ============= 6. 이미지 업로드 =============
add_title('6. 이미지 업로드', 1)

add_para('게시판 관리에서 이미지를 업로드하면 자동으로 URL이 발급됩니다.')
add_para('업로드 경로 형식:', bold=True)
add_code('/uploads/hospitals/{병원slug}/boards/{파일명}')

add_para('제한:', bold=True)
add_bullet('최대 10MB')
add_bullet('jpg, png, webp, gif, avif 지원')

add_para('파비콘은 SEO·병원 정보 → GEO 정보 → 파비콘 업로드에서 별도 등록.')

# ============= 7. SEO 설정 =============
add_title('7. SEO 설정', 1)

add_title('7.1 페이지별 SEO', 2)
add_para('SEO·병원 정보 → 페이지 SEO 탭에서 각 페이지의 메타태그를 설정합니다.')
add_bullet('Title 태그: 브라우저 탭에 표시되는 제목')
add_bullet('Meta Description: 검색 결과에 표시되는 설명')
add_bullet('OG 제목/설명/이미지: 카카오톡/페이스북 공유 시 미리보기')
add_bullet('Canonical URL: 정규 URL')

add_title('7.2 GEO 정보', 2)
add_para('병원의 위치/연락처/SNS 정보를 입력하면 자동으로 구조화 데이터(JSON-LD)에 반영됩니다.')

add_title('7.3 공통 코드', 2)
add_para('Google Tag Manager, 네이버 애널리틱스 등을 모든 페이지에 일괄 삽입.')
add_bullet('공통 메타 태그: <head>에 삽입')
add_bullet('공통 헤더 코드: <body> 시작 직후')
add_bullet('공통 본문 코드: 콘텐츠 영역 전')
add_bullet('공통 푸터 코드: </body> 직전')

# ============= 8. 자주 발생하는 오류 =============
add_title('8. 자주 발생하는 오류', 1)

add_title('8.1 JS가 작동하지 않을 때', 2)
add_bullet('JS 탭에 <script> 태그를 포함하지 마세요. (자동 제거되지만 권장하지 않음)')
add_bullet("DOMContentLoaded 이벤트는 자동 처리되지만, 불필요하면 제거 가능")
add_bullet('브라우저 콘솔(F12)에서 오류 확인')

add_title('8.2 모바일 메뉴가 메인 메뉴 뒤로 숨을 때', 2)
add_para('PC 드롭다운(.depth2)과 모바일 아코디언(.m-accordion .depth2)이 같은 클래스명이면 PC의 position:absolute가 모바일에도 적용됩니다.')
add_para('모바일 .depth2에 명시적으로 PC 스타일을 리셋하세요. (3.2 참고)')

add_title('8.3 폼 제출이 두 번 들어갈 때', 2)
add_para('블록 스니펫의 폼 JS는 중복 방지가 적용되어 있습니다. 직접 작성한 폼이라면 제출 후 disabled 처리하세요.')

add_title('8.4 페이지가 다른 병원 데이터로 보일 때', 2)
add_para('도메인 매핑이 잘못된 경우입니다. 슈퍼관리자에게 도메인 등록 확인을 요청하세요.')

add_title('8.5 한글 입력 시 글자가 중복되는 경우', 2)
add_para('관리자 폼은 한글 IME 처리가 적용되어 있습니다. 만약 발생하면 새로고침 후 재입력.')

# ============= 9. 미리보기와 배포 =============
add_title('9. 미리보기와 배포', 1)

add_title('9.1 도메인 연결 전 미리보기', 2)
add_para('서브도메인으로 실제 사이트 미리보기 가능:')
add_code('https://{병원slug}.metalink3.mycafe24.com')

add_para('예시: 슬러그가 wonju-dental인 병원')
add_code('https://wonju-dental.metalink3.mycafe24.com')

add_title('9.2 실제 도메인 연결', 2)
add_para('도메인 관리 메뉴에서 도메인을 등록한 후, DNS A 레코드를 서버 IP로 설정하면 즉시 적용됩니다.')

# ============= 10. 권장 작업 순서 =============
add_title('10. 권장 작업 순서', 1)

steps = [
    'SEO·병원 정보 → 병원 GEO 정보 입력 (병원명, 주소, 전화 등)',
    '파비콘 업로드',
    '헤더·푸터 → 공통 헤더 디자인 (반드시 모바일 메뉴 포함)',
    '헤더·푸터 → 공통 푸터 디자인',
    '메인 페이지 디자인',
    '서브 페이지 추가 및 디자인 (병원소개, 진료과목 등)',
    '게시판 그룹 생성 (필요 시)',
    '공지/이벤트 목록·상세 페이지 디자인',
    '팝업 설정 (이벤트 안내 등)',
    '페이지 SEO 메타태그 입력',
    '공통 코드 (GTM/애널리틱스) 등록',
    '서브도메인 미리보기로 전체 검수',
    '실제 도메인 연결',
]

for i, step in enumerate(steps, 1):
    add_para('{}. {}'.format(i, step))

# ============= 11. 지원 =============
add_title('11. 기술 지원', 1)
add_bullet('카카오톡 채널: https://pf.kakao.com/_xjxhLyn')
add_bullet('기술지원 게시판: https://mtlink.kr/TechSupport')

add_para('').paragraph_format.space_after = Pt(20)
footer = add_para('© METALINK CMS — 웹 퍼블리싱 디자이너 가이드')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in footer.runs:
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


import os
out_path = os.path.join(os.path.dirname(__file__), '..', 'docs', '웹퍼블리싱_디자이너_가이드.docx')
doc.save(out_path)
print('saved:', os.path.abspath(out_path))
